import os
import streamlit as st
import time
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders.word_document import Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_cohere import CohereEmbeddings, ChatCohere
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain import hub
from linkup import LinkupClient
import tempfile
from dotenv import load_dotenv
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing

# Load environment variables
load_dotenv()

# API Keys
cohere_key = os.getenv("COHERE_API_KEY")
qdrant_key = os.getenv("QDRANT_API_KEY")
qdrant_url = os.getenv("QDRANT_URL")
linkup_key = os.getenv("LINKUP_API_KEY")

# Initialize Linkup client
linkup_client = LinkupClient(api_key=linkup_key)

# Page configuration
st.set_page_config(
    page_title="VectorDB Search with LinkUp",
    page_icon="🔍",
    layout="wide"
)


st.markdown(
    """
    <h4 style="display: flex; align-items: center; justify-content: center;">
        <img src="https://avatars.githubusercontent.com/u/54850923?s=280&v=4" width="50" style="margin-right: 0px;"> 
        RAG Powered by 
        <img src="https://framerusercontent.com/images/wLLGrlJoyqYr9WvgZwzlw91A8U.png" width="120" style="margin: 0 7px;"> 
        & 
        <img src="https://sp-ao.shortpixel.ai/client/to_auto,q_lossy,ret_img/https://www.wpsolr.com/wp-content/uploads/elementor/thumbs/qdrant_logo_with_text-qmw293llp8vbuf66ezchw33iwsm60uhwo9ddqxx1q8.png" width="110" style="margin-left: 10px;">   
    </h4>
    """,
    unsafe_allow_html=True
)

# Session state initialization
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "chat_memory" not in st.session_state:
    st.session_state.chat_memory = []  # Separate memory for context
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "processed_file" not in st.session_state:
    st.session_state.processed_file = None
if "document_text" not in st.session_state:
    st.session_state.document_text = None
if "file_type" not in st.session_state:
    st.session_state.file_type = None
if "mode" not in st.session_state:
    st.session_state.mode = "Direct Chat"

# Initialize language models
embedding = CohereEmbeddings(model="embed-english-v3.0", cohere_api_key=cohere_key)
chat_model = ChatCohere(
    model="command-a-03-2025",
    temperature=0.4,
    verbose=True,
    streaming=True,
    cohere_api_key=cohere_key
)

# Qdrant Client Setup
client = QdrantClient(url=qdrant_url, api_key=qdrant_key, timeout=60)
COLLECTION_NAME = "new"

# Ensure the collection exists
try:
    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=1024, distance=Distance.COSINE),
    )
except Exception as e:
    if "already exists" not in str(e).lower():
        st.error(f"Error creating Qdrant collection: {e}")

# LinkUp search function
def perform_linkup_search(query, depth="standard", output_type="sourcedAnswer"):
    """
    Use Linkup API to search the web for information
    """
    try:
        response = linkup_client.search(
            query=query,
            depth=depth,
            output_type=output_type,
        )
        return response
    except Exception as e:
        st.error(f"Linkup search error: {e}")
        return {"error": str(e)}

# Format chat memory for context
def format_chat_memory_for_context(chat_memory, max_turns=5):
    context = ""
    # Take only the last few turns to avoid context length issues
    recent_memory = chat_memory[-max_turns*2:] if len(chat_memory) > max_turns*2 else chat_memory
    
    for message in recent_memory:
        role = message["role"].upper()
        content = message["content"]
        context += f"{role}: {content}\n\n"
    
    return context

# Document Processing Functions
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_txt(txt_path):
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        try:
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            st.error(f"Error reading text file: {e}")
            return ""
    except Exception as e:
        st.error(f"Error reading text file: {e}")
        return ""

def process_document(file, file_type):
    """Process uploaded document and return text chunks for vectorization"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_type}') as tmp_file:
        tmp_file.write(file.getvalue())
        tmp_path = tmp_file.name
    
    # Extract text based on file type
    if file_type == 'pdf':
        document_text = extract_text_from_pdf(tmp_path)
        loader = PyPDFLoader(tmp_path)
    elif file_type == 'docx':
        document_text = extract_text_from_docx(tmp_path)
        loader = Docx2txtLoader(tmp_path)
    elif file_type == 'txt':
        document_text = extract_text_from_txt(tmp_path)
        
        # For txt files, create a simple loader
        from langchain_core.documents import Document
        documents = [Document(page_content=document_text)]
        os.unlink(tmp_path)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, 
            chunk_overlap=200, 
            separators=["\n\n", "\n", " ", ""]
        )
        texts = text_splitter.split_documents(documents)
        return document_text, texts
    else:
        os.unlink(tmp_path)
        raise ValueError(f"Unsupported file type: {file_type}")
    
    # For PDF and DOCX, use langchain loaders
    documents = loader.load()
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, 
        chunk_overlap=200, 
        separators=["\n\n", "\n", " ", ""]
    )
    texts = text_splitter.split_documents(documents)
    os.unlink(tmp_path)
    
    return document_text, texts

def create_vector_store(texts):
    """Create or update the vector store with document chunks"""
    vector_store = QdrantVectorStore(
        client=client, 
        collection_name=COLLECTION_NAME, 
        embedding=embedding
    )
    vector_store.add_documents(texts)
    return vector_store

# Document preview function
def get_document_preview(file_path, file_extension):
    if file_extension == "pdf":
        try:
            doc = fitz.open(file_path)
            first_page = doc[0]
            pix = first_page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))
            img_bytes = pix.tobytes("png")
            doc.close()
            return img_bytes, "image"
        except:
            return "PDF Preview Unavailable", "text"
    elif file_extension == "docx":
        try:
            doc = docx.Document(file_path)
            preview_text = ""
            # Get first 5 paragraphs or fewer if document is smaller
            for i, para in enumerate(doc.paragraphs):
                if i < 5 and para.text.strip():
                    preview_text += para.text + "\n\n"
                if i >= 5:
                    preview_text += "..."
                    break
            return preview_text or "Document is empty or contains non-text content", "text"
        except:
            return "DOCX Preview Unavailable", "text"
    elif file_extension == "txt":
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                lines = file.readlines()
                # Get first 10 lines or fewer if file is smaller
                preview_lines = lines[:10]
                preview_text = "".join(preview_lines)
                if len(lines) > 10:
                    preview_text += "..."
                return preview_text or "Text file is empty", "text"
        except:
            return "TXT Preview Unavailable", "text"
    else:
        return "Preview not available for this file type", "text"

# Advanced query processing functions
def perform_vector_search(query, vectorstore):
    """Perform vector search and assess relevance of results"""
    if not vectorstore:
        return None, False
    
    retriever = vectorstore.as_retriever(
        search_type="mmr",  # Maximum marginal relevance
        search_kwargs={
            "k": 5,
            "fetch_k": 20,
            "lambda_mult": 0.5  # Balance between relevance and diversity
        }
    )
    
    # Retrieve documents
    retrieved_docs = retriever.get_relevant_documents(query)
    
    # Check if documents were found
    if not retrieved_docs:
        return None, False
    
    # Assess document relevance
    def assess_document_relevance(docs, query):
        if len(docs) >= 3:
            # If we found multiple documents, likely relevant
            return True
        
        # Simple keyword matching as fallback
        query_keywords = set(query.lower().split())
        matches = 0
        for doc in docs:
            doc_text = doc.page_content.lower()
            for keyword in query_keywords:
                if keyword in doc_text and len(keyword) > 3:  # Avoid matching short words
                    matches += 1
        
        # If we match at least half the keywords, consider relevant
        return matches >= len(query_keywords) / 2
    
    is_relevant = assess_document_relevance(retrieved_docs, query)
    return retrieved_docs, is_relevant

def process_query(query):
    """Process user query with intelligent routing logic"""
    # Get chat context
    chat_context = format_chat_memory_for_context(st.session_state.chat_memory)
    
    with st.spinner("Processing your question..."):
        # First stage: Check document relevance if we have a vector database
        doc_relevance = False
        retrieved_docs = None
        
        if st.session_state.vectorstore:
            retrieved_docs, doc_relevance = perform_vector_search(query, st.session_state.vectorstore)
        
        # Always assess if web search is needed, regardless of mode or document relevance
        assessment_prompt = f"""
        CHAT HISTORY:
        {chat_context}
        
        USER QUERY: {query}
        
        First, assess if this query requires up-to-date information, specific factual details, or 
        information that might be beyond your knowledge cutoff.
        
        Consider these factors:
        - Is the query about very recent events, current data, or trending topics?
        - Does it require specific factual information you might not have?
        - Would searching for information significantly improve the accuracy of your response?
        - Does it reference previous parts of our conversation?
        {f"- Note: I already have relevant document information, but determine if additional web search would provide value." if doc_relevance else ""}
        
        RESPOND ONLY WITH:
        "SEARCH_NEEDED" - If web search would significantly improve the response quality
        "DIRECT_ANSWER" - If you can confidently answer without additional information
        """
        
        # Use non-streaming for this internal prompt
        non_streaming_model = ChatCohere(
            model="command-a-03-2025",
            temperature=0.1,
            cohere_api_key=cohere_key
        )
        
        search_decision = non_streaming_model.invoke(assessment_prompt)
        
        # Based on assessment, perform search or answer directly
        if "SEARCH_NEEDED" in search_decision.content:
            with st.spinner("Searching for information..."):
                search_results = perform_linkup_search(query, depth="deep")
                
                # If we also have relevant document results, combine both sources
                if retrieved_docs:
                    # Create prompt with both document and search results
                    hybrid_prompt = f"""
                    CHAT HISTORY:
                    {chat_context}
                    
                    USER QUERY: {query}
                    
                    I have information from two sources to help answer this query:
                    
                    1. DOCUMENT CONTENT:
                    {' '.join([doc.page_content for doc in retrieved_docs[:3]])}
                    
                    2. WEB SEARCH RESULTS:
                    {search_results}
                    
                    Please provide a comprehensive answer that integrates information from both sources when relevant.
                    Make it clear which information comes from the document and which comes from web search.
                    Format your response clearly and consider our previous conversation when applicable.
                    """
                    
                    return {
                        "type": "hybrid_search",
                        "prompt": hybrid_prompt,
                        "doc_sources": retrieved_docs[:3],
                        "web_sources": search_results
                    }
                else:
                    # Only web search results are relevant
                    search_prompt = f"""
                    CHAT HISTORY:
                    {chat_context}
                    
                    USER QUERY: {query}
                    
                    Based on the following search results:
                    
                    {search_results}
                    
                    Please provide a comprehensive answer to the query. Remember to consider our previous conversation
                    when relevant. Format your response clearly and cite information from these search results when appropriate.
                    """
                    
                    return {
                        "type": "web_search",
                        "prompt": search_prompt,
                        "sources": search_results
                    }
        elif retrieved_docs:
            # Document results exist but web search not needed - use document only
            retrieval_qa_prompt = hub.pull("langchain-ai/retrieval-qa-chat")
            combine_docs_chain = create_stuff_documents_chain(chat_model, retrieval_qa_prompt)
            retrieval_chain = create_retrieval_chain(
                st.session_state.vectorstore.as_retriever(), 
                combine_docs_chain
            )
            
            return {
                "type": "document_rag",
                "chain": retrieval_chain,
                "query": f"""
                CHAT HISTORY:
                {chat_context}
                
                USER QUERY: {query}
                
                Answer based on the retrieved document content and our previous conversation.
                """,
                "sources": retrieved_docs
            }
        else:
            # Answer directly from model knowledge - no relevant docs or web search needed
            direct_prompt = f"""
            CHAT HISTORY:
            {chat_context}
            
            USER QUERY: {query}
            
            Provide a helpful, accurate response based on your knowledge and our previous conversation.
            If you're uncertain about any aspect, mention this uncertainty.
            """
            
            return {
                "type": "direct_answer",
                "prompt": direct_prompt
            }

# Sidebar

st.sidebar.image("Hero-Image.jpg", use_container_width=True)


# Mode selection
st.sidebar.header("📤 Upload & Mode Selection")
mode_options = ["💬 GeneralChat", "📄 Document Analysis"]
selected_mode = st.sidebar.radio("Select mode", mode_options)

if selected_mode == "💬 GeneralChat":
    st.session_state.mode = "Direct Chat"
elif selected_mode == "📄 Document Analysis":
    st.session_state.mode = "Document Analysis"
    
    # Document uploader
    uploaded_file = st.sidebar.file_uploader("Upload Document", type=["pdf", "txt", "docx"])
    
    if uploaded_file:
        # Check if we need to process a new file
        if st.session_state.get('last_file_name') != uploaded_file.name:
            st.session_state.chat_history = []
            st.session_state.chat_memory = []
            
            with st.spinner("Processing document..."):
                file_type = uploaded_file.name.split('.')[-1].lower()
                st.session_state['last_file_name'] = uploaded_file.name
                
                # Process document and create vector store
                document_text, text_chunks = process_document(uploaded_file, file_type)
                st.session_state.document_text = document_text
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_type}') as temp_file:
                    temp_file.write(uploaded_file.getvalue())
                    file_path = temp_file.name
                
                st.session_state.processed_file = file_path
                st.session_state.vectorstore = create_vector_store(text_chunks)
                
                # Display document preview
                preview_content, preview_type = get_document_preview(file_path, file_type)
                if preview_type == "image":
                    st.sidebar.image(preview_content, caption=f"{file_type.upper()} Preview", use_container_width=True)
                else:
                    st.sidebar.text_area("Document Preview", preview_content, height=200, disabled=True)
                
                # Display document info
                text_length = len(document_text)
                st.sidebar.info(f"{file_type.upper()} processed: {uploaded_file.name} ({text_length} characters)")
                st.sidebar.success(f"Document processed successfully!")

# Sidebar controls
st.sidebar.header("⚙️ Controls")
if st.sidebar.button('🧹 Clear Chat History'):
    st.session_state.chat_history = []
    st.session_state.chat_memory = []
    st.rerun()
if st.sidebar.button('🗑️ Clear All Data'):
    try:
        client.delete_collection(COLLECTION_NAME)
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=1024, distance=Distance.COSINE),
        )
    except:
        pass
    st.session_state.vectorstore = None
    st.session_state.processed_file = None
    st.session_state.document_text = None
    st.session_state.chat_history = []
    st.session_state.chat_memory = []
    if 'last_file_name' in st.session_state:
        del st.session_state['last_file_name']
    st.sidebar.success("All data cleared!")
    st.rerun()

# Display chat history
for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Chat input
input_placeholder = "Ask any question..." if st.session_state.mode == "Direct Chat" else "Ask a question about your document..."
if query := st.chat_input(input_placeholder):
    # Add user query to chat history
    st.session_state.chat_history.append({"role": "user", "content": query})
    st.session_state.chat_memory.append({"role": "user", "content": query})
    
    with st.chat_message("user"):
        st.markdown(query)
    
    with st.chat_message("assistant"):
        # Process the query
        result = process_query(query)
        
        # Response placeholder for streaming
        response_placeholder = st.empty()
        full_response = ""
        
        # Handle different response types
        if result["type"] == "document_rag":
            # Stream the response for document-based answers
            for chunk in result["chain"].stream({"input": result["query"]}):
                if "answer" in chunk:
                    full_response += chunk["answer"]
                    response_placeholder.markdown(full_response)
            
            # Show source information
            with st.expander(f"🔍 Document Sources", expanded=False):
                for i, doc in enumerate(result['sources'][:3]):  # Limit to first 3 sources
                    st.markdown(f"**Source {i+1}**\n{doc.page_content[:300]}...")
                    
        elif result["type"] == "web_search":
            # Stream with web search results
            for chunk in chat_model.stream(result["prompt"]):
                full_response += chunk.content
                response_placeholder.markdown(full_response)
            
            # Show source information
            with st.expander(f"🔍 Web Search Sources", expanded=False):
                st.markdown("This response was generated using Linkup search results")
                
        elif result["type"] == "hybrid_search":
            # Stream with combined document and web search results
            for chunk in chat_model.stream(result["prompt"]):
                full_response += chunk.content
                response_placeholder.markdown(full_response)
            
            # Show source information
            with st.expander(f"🔍 Combined Sources", expanded=False):
                st.markdown("### Document Sources")
                for i, doc in enumerate(result['doc_sources']):
                    st.markdown(f"**Document Source {i+1}**\n{doc.page_content[:300]}...")
                
                st.markdown("### Web Search Sources")
                st.markdown("This response also used Linkup search results to provide additional context")
                
        else:  # direct_answer
            # Stream direct model answer
            for chunk in chat_model.stream(result["prompt"]):
                full_response += chunk.content
                response_placeholder.markdown(full_response)
        
        # Add to chat history
        st.session_state.chat_history.append({
            "role": "assistant", 
            "content": full_response
        })
        st.session_state.chat_memory.append({
            "role": "assistant", 
            "content": full_response
        })
